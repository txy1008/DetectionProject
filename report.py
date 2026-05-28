"""
report.py — 检测报告 Word 文档生成
生成 .docx 报告，包含：报告时间、检测依据、目标统计汇总、CSV 数据表、饼图与柱状图
"""

import os
import io
from datetime import datetime
from collections import Counter

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config import RESULTS_DIR
from database import db_get_records

# ── 中文字体设置 ─────────────────────────────────────────────────────────────
_CN_FONT = None
for fn in ["SimHei", "Microsoft YaHei", "SimSun", "DengXian"]:
    fp = font_manager.findfont(fn, fallback_to_default=False)
    if fp and os.path.exists(fp):
        _CN_FONT = fn
        break
if _CN_FONT:
    plt.rcParams["font.sans-serif"] = [_CN_FONT]
plt.rcParams["axes.unicode_minus"] = False

# 类别映射
MC = ["person"]
NC = ["car", "truck", "bus", "motorcycle", "bicycle", "train", "boat", "airplane"]
PC = ["bird", "cat", "dog", "horse", "sheep", "cow", "elephant", "bear", "zebra", "giraffe"]
CLR = ["#f48c42", "#4fc3f7", "#81c784", "#ce93d8", "#80cbc4", "#ffcc80", "#ef9a9a"]


def _cat(cls_name: str) -> str:
    if cls_name in MC:
        return "人员"
    elif cls_name in NC:
        return "车辆"
    elif cls_name in PC:
        return "动物"
    return "其他"


def _make_pie_chart(class_counts: dict) -> bytes:
    """生成饼图，返回 PNG bytes"""
    labels = list(class_counts.keys())
    values = list(class_counts.values())
    colors = [CLR[i % len(CLR)] for i in range(len(labels))]

    fig, ax = plt.subplots(figsize=(5, 4))
    ax.pie(values, labels=labels, colors=colors, autopct="%1.1f%%", startangle=90,
           textprops={"fontsize": 10})
    ax.set_title("目标类别占比", fontsize=13, fontweight="bold")
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _make_bar_chart(class_counts: dict) -> bytes:
    """生成柱状图，返回 PNG bytes"""
    labels = list(class_counts.keys())
    values = list(class_counts.values())
    colors = [CLR[i % len(CLR)] for i in range(len(labels))]

    fig, ax = plt.subplots(figsize=(5, 4))
    bars = ax.bar(labels, values, color=colors)
    ax.set_title("目标数量统计", fontsize=13, fontweight="bold")
    ax.set_ylabel("数量")
    ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
    for bar, v in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.2,
                str(v), ha="center", fontsize=10, fontweight="bold")
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10):
    """设置表格单元格文本"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)


def generate_report(session_name: str) -> str:
    """
    根据 session_name 生成 Word 报告，保存在 results/ 目录。
    返回生成的 docx 文件路径。
    """
    records = db_get_records(session_name=session_name)
    if not records:
        raise ValueError(f"会话 '{session_name}' 没有检测记录")

    now = datetime.now()
    doc = Document()

    # ── 标题 ──
    title = doc.add_heading("目标检测报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 报告信息 ──
    doc.add_heading("一、报告信息", level=1)
    info_tbl = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("报告名称", session_name),
        ("报告时间", now.strftime("%Y年%m月%d日 %H:%M:%S")),
        ("检测模型", "YOLOv8n + BotSORT 多目标追踪"),
        ("检测依据", "基于 COCO 数据集预训练权重，置信度阈值过滤"),
    ]
    for i, (k, v) in enumerate(info_data):
        _set_cell(info_tbl.rows[i].cells[0], k, bold=True, size=10)
        _set_cell(info_tbl.rows[i].cells[1], v, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)

    doc.add_paragraph("")

    # ── 检测汇总 ──
    doc.add_heading("二、检测结果汇总", level=1)

    class_counts = Counter(r["class_name"] for r in records)
    total = len(records)
    m_cnt = sum(1 for r in records if r["class_name"] in MC)
    n_cnt = sum(1 for r in records if r["class_name"] in NC)
    p_cnt = sum(1 for r in records if r["class_name"] in PC)

    summary_p = doc.add_paragraph()
    summary_p.add_run(f"本次检测共发现 ").font.size = Pt(11)
    run_total = summary_p.add_run(f"{total}")
    run_total.bold = True
    run_total.font.size = Pt(14)
    run_total.font.color.rgb = RGBColor(0x22, 0x7C, 0xB6)
    summary_p.add_run(f" 个目标，其中：").font.size = Pt(11)

    cat_tbl = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
    cat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (cat, cnt) in enumerate([("类别", "数量"), ("🧑 人员", m_cnt), ("🚗 车辆", n_cnt), ("🐾 动物", p_cnt)]):
        pass
    headers = ["类别", "人员 (person)", "车辆 (car/truck/bus…)", "动物 (dog/cat/bird…)"]
    values = ["数量", str(m_cnt), str(n_cnt), str(p_cnt)]
    for i, h in enumerate(headers):
        _set_cell(cat_tbl.rows[0].cells[i], h, bold=True)
    row = cat_tbl.add_row()
    for i, v in enumerate(values):
        _set_cell(row.cells[i], v, bold=(i == 0))

    # 细分类别
    doc.add_paragraph("")
    doc.add_paragraph("各类别细分统计：", style="List Bullet")
    for cls, cnt in class_counts.most_common():
        doc.add_paragraph(f"{cls} ({_cat(cls)}): {cnt} 个", style="List Bullet 2")

    doc.add_paragraph("")

    # ── 数据分析图表 ──
    doc.add_heading("三、数据分析图表", level=1)

    pie_bytes = _make_pie_chart(dict(class_counts))
    bar_bytes = _make_bar_chart(dict(class_counts))

    doc.add_paragraph("1. 目标类别占比（饼图）：")
    doc.add_picture(io.BytesIO(pie_bytes), width=Inches(4.5))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("2. 目标数量统计（柱状图）：")
    doc.add_picture(io.BytesIO(bar_bytes), width=Inches(4.5))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # ── CSV 数据表格 ──
    doc.add_heading("四、检测明细记录", level=1)
    doc.add_paragraph(f"共 {total} 条记录，按检测时间排列：")

    col_headers = ["序号", "目标ID", "类别", "检测时间", "来源"]
    tbl = doc.add_table(rows=1, cols=len(col_headers), style="Light Grid Accent 1")
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(col_headers):
        _set_cell(tbl.rows[0].cells[i], h, bold=True)

    for idx, rec in enumerate(records, 1):
        row = tbl.add_row()
        _set_cell(row.cells[0], str(idx))
        _set_cell(row.cells[1], str(rec.get("obj_id", "")))
        _set_cell(row.cells[2], rec.get("class_name", ""))
        _set_cell(row.cells[3], rec.get("detect_time", ""))
        _set_cell(row.cells[4], rec.get("source_type", ""))

    doc.add_paragraph("")

    # ── 页脚备注 ──
    doc.add_paragraph("").add_run("— 报告结束 —").italic = True
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run(f"由 YOLOv8 智慧路口视频监控系统自动生成 · {now.strftime('%Y-%m-%d %H:%M')}")
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── 保存 ──
    docx_path = os.path.join(RESULTS_DIR, f"{session_name}.docx")
    doc.save(docx_path)
    return docx_path
