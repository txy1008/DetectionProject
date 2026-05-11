import os
import io
import pandas as pd
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
from PySide6.QtWidgets import (QDialog, QVBoxLayout, QHBoxLayout, QPushButton,
                               QComboBox, QLabel, QFileDialog, QMessageBox)
from PySide6.QtCore import Qt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import config


class HistoryDialog(QDialog):
    """历史数据分析对话框"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("📊 历史数据分析")
        self.resize(900, 650)
        self.setStyleSheet("background-color: #f8f9fa;")

        layout = QVBoxLayout(self)

        # --- 顶部控制栏 ---
        top_bar = QHBoxLayout()

        top_bar.addWidget(QLabel("选择场次:"))
        self.combo_session = QComboBox()
        self.combo_session.setMinimumWidth(250)
        self._load_sessions()
        top_bar.addWidget(self.combo_session)

        btn_analyze = QPushButton("📈 分析")
        btn_analyze.setStyleSheet("padding: 8px 16px; background-color: #0078D7; color: white; border-radius: 4px;")
        btn_analyze.clicked.connect(self.analyze)
        top_bar.addWidget(btn_analyze)

        btn_export = QPushButton("📄 导出 Word 报告")
        btn_export.setStyleSheet("padding: 8px 16px; background-color: #28a745; color: white; border-radius: 4px;")
        btn_export.clicked.connect(self.export_word)
        top_bar.addWidget(btn_export)

        top_bar.addStretch()
        layout.addLayout(top_bar)

        # --- 图表区域 ---
        self.figure, self.axes = plt.subplots(1, 2, figsize=(10, 4))
        self.figure.set_facecolor("#f8f9fa")
        self.canvas = FigureCanvas(self.figure)
        layout.addWidget(self.canvas)

        # --- 摘要标签 ---
        self.summary_label = QLabel("选择场次后点击「分析」查看数据")
        self.summary_label.setAlignment(Qt.AlignCenter)
        self.summary_label.setStyleSheet("font-size: 14px; color: #333; padding: 10px; background-color: white; border-radius: 6px;")
        layout.addWidget(self.summary_label)

        self.current_df = None

    def _load_sessions(self):
        """加载所有可用的 CSV 报告"""
        self.combo_session.clear()
        if not os.path.exists(config.RESULTS_DIR):
            return
        csv_files = [f for f in os.listdir(config.RESULTS_DIR) if f.endswith(".csv")]
        csv_files.sort(reverse=True)
        for f in csv_files:
            self.combo_session.addItem(f)

    def analyze(self):
        """分析选中的场次数据"""
        filename = self.combo_session.currentText()
        if not filename:
            return

        filepath = os.path.join(config.RESULTS_DIR, filename)
        try:
            df = pd.read_csv(filepath, encoding='utf-8-sig')
        except Exception as e:
            self.summary_label.setText(f"读取失败: {e}")
            return

        if df.empty or "类别" not in df.columns:
            self.summary_label.setText("数据为空或格式不匹配")
            return

        self.current_df = df
        self._draw_charts(df, filename)

    def _draw_charts(self, df, title):
        """绘制分析图表"""
        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False

        for ax in self.axes:
            ax.clear()

        # 图1: 饼图 - 各类别占比
        category_counts = df["类别"].value_counts()
        colors = {"car": "#e67e22", "bicycle": "#3498db", "person": "#2ecc71"}
        pie_colors = [colors.get(c, "#95a5a6") for c in category_counts.index]
        labels = {"car": "机动车", "bicycle": "非机动车", "person": "行人"}
        pie_labels = [labels.get(c, c) for c in category_counts.index]

        self.axes[0].pie(category_counts.values, labels=pie_labels, colors=pie_colors,
                         autopct='%1.1f%%', startangle=90, textprops={'fontsize': 11})
        self.axes[0].set_title("目标类别占比", fontsize=13, fontweight='bold')

        # 图2: 柱状图 - 各类别数量
        bars = self.axes[1].bar(pie_labels, category_counts.values, color=pie_colors, edgecolor='white')
        self.axes[1].set_title("目标数量统计", fontsize=13, fontweight='bold')
        self.axes[1].set_ylabel("数量")
        for bar, val in zip(bars, category_counts.values):
            self.axes[1].text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3,
                              str(val), ha='center', fontsize=12, fontweight='bold')

        self.figure.suptitle(title.replace('.csv', ''), fontsize=14, fontweight='bold')
        self.figure.tight_layout()
        self.canvas.draw()

        # 更新摘要
        total = len(df)
        car_n = len(df[df["类别"] == "car"])
        bike_n = len(df[df["类别"] == "bicycle"])
        person_n = len(df[df["类别"] == "person"])
        self.summary_label.setText(
            f"总检测数: {total}  |  🚗 机动车: {car_n}  |  🚲 非机动车: {bike_n}  |  🚶 行人: {person_n}")

    def export_word(self):
        """导出当前分析结果为 Word 文档（表格 + 图表合并）"""
        if self.current_df is None:
            QMessageBox.warning(self, "提示", "请先分析一个场次的数据")
            return

        session_name = self.combo_session.currentText().replace('.csv', '')
        default_path = os.path.join(config.RESULTS_DIR, f"{session_name}.docx")
        path, _ = QFileDialog.getSaveFileName(self, "保存 Word 报告", default_path, "Word Files (*.docx)")
        if not path:
            return

        try:
            df = self.current_df
            doc = Document()

            # 标题
            title = doc.add_heading('智慧路口视频监控系统 — 检测报告', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'场次: {session_name}')
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(100, 100, 100)
            doc.add_paragraph()

            # 一、统计摘要
            doc.add_heading('一、统计摘要', level=1)
            total = len(df)
            car_n = len(df[df['类别'] == 'car'])
            bike_n = len(df[df['类别'] == 'bicycle'])
            person_n = len(df[df['类别'] == 'person'])
            t_sum = doc.add_table(rows=3, cols=2, style='Light Shading Accent 1')
            t_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, (k, v) in enumerate([
                ('检测总数', str(total)),
                ('机动车 / 非机动车 / 行人', f'{car_n} / {bike_n} / {person_n}'),
                ('数据来源', session_name),
            ]):
                t_sum.cell(i, 0).text = k
                t_sum.cell(i, 1).text = v
            doc.add_paragraph()

            # 二、检测记录明细
            doc.add_heading('二、检测记录明细', level=1)
            headers = list(df.columns)
            table = doc.add_table(rows=1, cols=len(headers), style='Light Grid Accent 1')
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.bold = True
                        r.font.size = Pt(9)
            for _, row in df.iterrows():
                cells = table.add_row().cells
                for j, h in enumerate(headers):
                    cells[j].text = str(row[h])
                    for para in cells[j].paragraphs:
                        for r in para.runs:
                            r.font.size = Pt(8)
            doc.add_paragraph()

            # 三、数据分析图表
            doc.add_heading('三、数据分析图表', level=1)

            # 饼图
            fig1, ax1 = plt.subplots(figsize=(5, 4))
            category_counts = df['类别'].value_counts()
            colors_map = {'car': '#e67e22', 'bicycle': '#3498db', 'person': '#2ecc71'}
            labels_map = {'car': '机动车', 'bicycle': '非机动车', 'person': '行人'}
            pie_colors = [colors_map.get(c, '#95a5a6') for c in category_counts.index]
            pie_labels = [labels_map.get(c, c) for c in category_counts.index]
            ax1.pie(category_counts.values, labels=pie_labels, colors=pie_colors,
                    autopct='%1.1f%%', startangle=90, textprops={'fontsize': 11})
            ax1.set_title('目标类别占比', fontsize=13, fontweight='bold')
            fig1.tight_layout()
            buf1 = io.BytesIO()
            fig1.savefig(buf1, format='png', dpi=150)
            buf1.seek(0)
            plt.close(fig1)
            p1 = doc.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.add_run().add_picture(buf1, width=Inches(4.5))

            # 柱状图
            fig2, ax2 = plt.subplots(figsize=(5, 4))
            bars = ax2.bar(pie_labels, category_counts.values, color=pie_colors, edgecolor='white')
            ax2.set_title('目标数量统计', fontsize=13, fontweight='bold')
            ax2.set_ylabel('数量')
            for bar, val in zip(bars, category_counts.values):
                ax2.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3,
                         str(val), ha='center', fontsize=12, fontweight='bold')
            fig2.tight_layout()
            buf2 = io.BytesIO()
            fig2.savefig(buf2, format='png', dpi=150)
            buf2.seek(0)
            plt.close(fig2)
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.add_run().add_picture(buf2, width=Inches(4.5))

            doc.save(path)
            QMessageBox.information(self, "导出成功", f"Word 报告已保存至:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "导出失败", f"生成 Word 失败: {e}")
