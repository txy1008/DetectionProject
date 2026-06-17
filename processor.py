import cv2
import os
import io
import numpy as np
import mysql.connector
import pandas as pd
import torch
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from datetime import datetime
from ultralytics import YOLO
from deep_sort_realtime.deepsort_tracker import DeepSort
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import config


class DetectionProcessor:
    def __init__(self):
        self.model = YOLO(config.MODEL_PATH)
        self.device = config.YOLO_DEVICE if torch.cuda.is_available() else "cpu"
        self.model.to(self.device)
        self.confirmed_ids = set()
        self.id_buffer = {}
        self.session_records = []

        self.save_count = 0
        self.counters = {"car": 0, "bicycle": 0, "person": 0}
        self.db_connected = True

        self.db_config = config.DB_CONFIG
        self.class_mapping = self._build_class_mapping()
        self.target_classes = list(self.class_mapping.keys())
        self.session_name = ""
        self.session_path = ""
        self.session_user = "default"

        # DeepSORT 追踪器
        self.tracker = self._create_tracker()

        # 轨迹历史（用于速度估计）
        self.track_history = {}   # {id: [(cx, cy), ...]}

        # 动态置信度阈值（UI 可调）
        self.conf_threshold = config.CONF_VIDEO

        # 速度估计相关 (像素/帧)
        self.speed_estimates = {}  # {id: speed_px_per_frame}
        self.fps_estimate = 15    # 默认 FPS——处理帧率

    @staticmethod
    def _create_tracker():
        """创建 DeepSORT 追踪器实例"""
        return DeepSort(
            max_age=config.DEEPSORT_MAX_AGE,
            n_init=config.DEEPSORT_N_INIT,
            nn_budget=config.DEEPSORT_NN_BUDGET,
            embedder=config.DEEPSORT_EMBEDDER,
            embedder_gpu=config.DEEPSORT_EMBEDDER_GPU
        )

    def _build_class_mapping(self):
        mapping = {}
        aliases = {
            "person": "person",
            "pedestrian": "person",
            "car": "car",
            "bus": "car",
            "truck": "car",
            "bicycle": "bicycle",
            "motorcycle": "bicycle",
            "motorbike": "bicycle",
            "bike": "bicycle",
        }
        for cls_id, name in self.model.names.items():
            key = str(name).lower().strip()
            if key in aliases:
                mapping[int(cls_id)] = aliases[key]
        return mapping or config.CLASS_MAPPING

    def change_model(self, new_model_path):
        """动态切换YOLO模型"""
        try:
            print(f"正在加载新模型: {new_model_path}")
            self.model = YOLO(new_model_path)
            self.device = config.YOLO_DEVICE if torch.cuda.is_available() else "cpu"
            self.model.to(self.device)
            self.class_mapping = self._build_class_mapping()
            self.target_classes = list(self.class_mapping.keys())
            # 同步更新 config.MODEL_PATH
            config.MODEL_PATH = new_model_path
            return True
        except Exception as e:
            print(f"模型加载失败: {e}")
            return False

    def start_session(self, username: str = "default"):
        """初始化新场次"""
        self.session_user = username
        today = datetime.now().strftime("%Y%m%d")
        
        user_results_dir = os.path.join(config.RESULTS_DIR, username)
        user_captures_dir = os.path.join(config.CAPTURES_DIR, username)
        
        os.makedirs(user_results_dir, exist_ok=True)
        os.makedirs(user_captures_dir, exist_ok=True)

        existing_reports = [f for f in os.listdir(user_results_dir)
                            if f.startswith(f"报告_{today}") and f.endswith(".csv")]
        run_index = len(existing_reports) + 1

        self.session_name = f"报告_{today}_第{run_index}次"
        self.session_path = os.path.abspath(f"{user_captures_dir}/{self.session_name}")
        os.makedirs(self.session_path, exist_ok=True)

        self.confirmed_ids.clear()
        self.id_buffer.clear()
        self.session_records = []
        self.save_count = 0
        self.counters = {"car": 0, "bicycle": 0, "person": 0}
        self.track_history.clear()
        self.speed_estimates.clear()
        self.tracker = self._create_tracker()  # 重置 DeepSORT 追踪器
        self.db_connected = True
        
        pd.DataFrame(columns=["ID", "类别", "时间", "存储路径"]).to_csv(
            os.path.join(user_results_dir, f"{self.session_name}.csv"), index=False, encoding='utf-8-sig')
        return self.session_name

    def process_frame(self, frame, is_image=False):
        """处理图像或视频帧"""
        h, w = frame.shape[:2]

        # ========== YOLO 检测（不做追踪，追踪交给 DeepSORT） ==========
        conf_threshold = self.conf_threshold
        results = self.model(frame, classes=self.target_classes, conf=conf_threshold,
                             device=self.device,
                             iou=config.IOU_THRESHOLD, verbose=False)

        annotated_frame = frame.copy()
        new_records = []

        if results[0].boxes is not None and len(results[0].boxes) > 0:
            boxes = results[0].boxes.xyxy.cpu().numpy()
            clss = results[0].boxes.cls.cpu().numpy().astype(int)
            confs = results[0].boxes.conf.cpu().numpy()

            if is_image:
                # 单图模式：不追踪
                self._last_image_detections = []
                for box, cls, conf in zip(boxes, clss, confs):
                    obj_name = self.class_mapping.get(cls, "unknown")

                    self.save_count += 1
                    if obj_name in self.counters:
                        self.counters[obj_name] += 1
                    display_id = self.save_count
                    record = self.save_data(frame, box, display_id, obj_name, conf, self.save_count)
                    new_records.append(record)

                    color_map = {"person": (0, 255, 0), "car": (255, 128, 0), "bicycle": (255, 200, 0)}
                    color = color_map.get(obj_name, (128, 128, 128))
                    cv2.rectangle(annotated_frame, (int(box[0]), int(box[1])), (int(box[2]), int(box[3])), color, 2)
                    cv2.putText(annotated_frame, f"{obj_name} #{display_id} {conf:.2f}",
                                (int(box[0]), int(box[1] - 10)), cv2.FONT_HERSHEY_SIMPLEX, 0.5, color, 2)

                    # 缓存检测详情供 UI 筛选
                    self._last_image_detections.append({
                        'id': display_id, 'box': box.tolist(),
                        'name': obj_name, 'conf': float(conf), 'color': color
                    })
            else:
                # ========== 视频模式：DeepSORT 追踪 ==========
                # 准备 DeepSORT 输入格式: ([left, top, w, h], confidence, class_id)
                raw_detections = []
                for box, cls, conf in zip(boxes, clss, confs):
                    x1, y1, x2, y2 = box
                    bw, bh = x2 - x1, y2 - y1
                    raw_detections.append(([x1, y1, bw, bh], float(conf), int(cls)))

                # DeepSORT 更新（传入原图用于提取外观特征），加入异常保护以防特征维数冲突崩溃
                try:
                    tracks = self.tracker.update_tracks(raw_detections, frame=frame)
                except Exception as e:
                    print(f"DeepSORT tracking conflict detected ({e}), resetting tracker gracefully to prevent crash.")
                    self.tracker = self._create_tracker()
                    tracks = []

                for track in tracks:
                    # 只处理当前帧匹配到的轨迹，过滤掉“幽灵”轨迹
                    if track.time_since_update > 0:
                        continue

                    track_id = track.track_id
                    ltrb = track.to_ltrb()  # [x1, y1, x2, y2]
                    det_class = track.det_class
                    det_conf = track.det_conf if track.det_conf is not None else 0.0
                    obj_name = self.class_mapping.get(det_class, "unknown")
                    is_confirmed = track.is_confirmed()

                    x1, y1, x2, y2 = ltrb

                    # 计算目标中心点
                    cx = int((x1 + x2) / 2)
                    cy = int((y1 + y2) / 2)

                    # 轨迹记录
                    if track_id not in self.track_history:
                        self.track_history[track_id] = []
                    self.track_history[track_id].append((cx, cy))

                    # 速度估计
                    hist = self.track_history[track_id]
                    if len(hist) >= 5:
                        dx = hist[-1][0] - hist[-5][0]
                        dy = hist[-1][1] - hist[-5][1]
                        dist = (dx ** 2 + dy ** 2) ** 0.5
                        self.speed_estimates[track_id] = dist / 5

                    # 保存新目标（仅已确认轨迹）
                    if not is_confirmed:
                        pass  # 未确认轨迹只画框，不计数
                    self.id_buffer[track_id] = self.id_buffer.get(track_id, 0) + 1
                    if is_confirmed and self.id_buffer[track_id] >= config.STABLE_FRAMES and track_id not in self.confirmed_ids:
                        self.confirmed_ids.add(track_id)
                        self.save_count += 1
                        if obj_name in self.counters:
                            self.counters[obj_name] += 1
                        record = self.save_data(frame, [x1, y1, x2, y2], track_id, obj_name, det_conf, self.save_count)
                        new_records.append(record)

                    # 绘制检测框
                    color_map = {"person": (0, 255, 0), "car": (255, 128, 0), "bicycle": (255, 200, 0)}
                    color = color_map.get(obj_name, (128, 128, 128))
                    cv2.rectangle(annotated_frame, (int(x1), int(y1)), (int(x2), int(y2)), color, 2)
                    label = f"{obj_name} ID:{track_id} {det_conf:.2f}"
                    cv2.putText(annotated_frame, label, (int(x1), int(y1 - 10)),
                                cv2.FONT_HERSHEY_SIMPLEX, 0.5, color, 2)

        return annotated_frame, new_records

    def save_data(self, frame, box, display_id, obj_name, conf, seq_num):
        """保存截图并存库"""
        cat_dir = os.path.join(self.session_path, obj_name)
        os.makedirs(cat_dir, exist_ok=True)

        now = datetime.now()
        # 关键修改：文件名中加入 seq_num 序号，防止单秒内多个目标互相覆盖
        img_name = f"{obj_name}_ID{display_id}_No{seq_num}_{now.strftime('%H%M%S')}.jpg"
        full_path = os.path.abspath(os.path.join(cat_dir, img_name))

        x1, y1, x2, y2 = map(int, box)
        crop = frame[max(0, y1):y2, max(0, x1):x2]
        cv2.imwrite(full_path, crop)

        try:
            conn = mysql.connector.connect(**self.db_config)
            cursor = conn.cursor()
            sql = "INSERT INTO detections (session_id, obj_id, category, img_name, img_path, confidence, detect_time) VALUES (%s, %s, %s, %s, %s, %s, %s)"
            db_id = int(display_id) if isinstance(display_id, int) else 0
            cursor.execute(sql, (self.session_name, db_id, obj_name, img_name, full_path, float(conf), now))
            conn.commit()
            cursor.close()
            conn.close()
        except Exception as e:
            self.db_connected = False
            print(f"[数据库] 写入失败: {e}")

        rec = {"ID": display_id, "类别": obj_name, "时间": now.strftime("%H:%M:%S"), "存储路径": full_path}
        self.session_records.append(rec)
        user_results_dir = os.path.join(config.RESULTS_DIR, self.session_user)
        os.makedirs(user_results_dir, exist_ok=True)
        pd.DataFrame(self.session_records).to_csv(
            os.path.join(user_results_dir, f"{self.session_name}.csv"), index=False, encoding='utf-8-sig')

        return {
            "id": display_id,
            "type": obj_name,
            "category": obj_name,
            "confidence": float(conf),
            "time": rec["时间"],
            "img": img_name,
            "path": full_path,
            "box": [float(x1), float(y1), float(x2), float(y2)]
        }

    def get_stats(self):
        """返回当前场次的实时统计数据"""
        return {**self.counters}

    def get_summary(self):
        """返回场次结束时的统计摘要"""
        total = sum(self.counters.values())
        return {
            "session": self.session_name,
            "total": total,
            "car": self.counters["car"],
            "bicycle": self.counters["bicycle"],
            "person": self.counters["person"],
            "db_status": "已连接" if self.db_connected else "未连接"
        }

    def generate_word_report(self, save_path=None):
        """生成 Word 报告：合并检测记录表格 + 统计图表"""
        if not self.session_records:
            return None

        doc = Document()

        # ===== 标题 =====
        title = doc.add_heading('检测报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'场次: {self.session_name}    生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()

        # ===== 一、统计摘要 =====
        doc.add_heading('一、统计摘要', level=1)
        summary = self.get_summary()
        total = summary['total']
        table_summary = doc.add_table(rows=3, cols=2, style='Light Shading Accent 1')
        table_summary.alignment = WD_TABLE_ALIGNMENT.CENTER
        summary_data = [
            ('检测总数', str(total)),
            ('机动车 / 非机动车 / 行人', f"{summary['car']} / {summary['bicycle']} / {summary['person']}"),
            ('数据库状态', summary['db_status']),
        ]
        for i, (key, val) in enumerate(summary_data):
            table_summary.cell(i, 0).text = key
            table_summary.cell(i, 1).text = val

        doc.add_paragraph()

        # ===== 二、检测记录明细 =====
        doc.add_heading('二、检测记录明细', level=1)
        df = pd.DataFrame(self.session_records)
        headers = list(df.columns)
        table = doc.add_table(rows=1, cols=len(headers), style='Light Grid Accent 1')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # 表体
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for j, h in enumerate(headers):
                cells[j].text = str(row[h])
                for paragraph in cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        doc.add_paragraph()

        # ===== 三、数据分析图表 =====
        doc.add_heading('三、数据分析图表', level=1)

        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False

        category_counts = df['类别'].value_counts()
        colors_map = {'car': '#e67e22', 'bicycle': '#3498db', 'person': '#2ecc71'}
        labels_map = {'car': '机动车', 'bicycle': '非机动车', 'person': '行人'}
        pie_colors = [colors_map.get(c, '#95a5a6') for c in category_counts.index]
        pie_labels = [labels_map.get(c, c) for c in category_counts.index]

        # 图1: 饼图
        fig1, ax1 = plt.subplots(figsize=(5, 4))
        ax1.pie(category_counts.values, labels=pie_labels, colors=pie_colors,
                autopct='%1.1f%%', startangle=90, textprops={'fontsize': 11})
        ax1.set_title('目标类别占比', fontsize=13, fontweight='bold')
        fig1.tight_layout()

        buf1 = io.BytesIO()
        fig1.savefig(buf1, format='png', dpi=150)
        buf1.seek(0)
        plt.close(fig1)

        p_chart1 = doc.add_paragraph()
        p_chart1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img1 = p_chart1.add_run()
        run_img1.add_picture(buf1, width=Inches(4.5))

        # 图2: 柱状图
        fig2, ax2 = plt.subplots(figsize=(5, 4))
        bars = ax2.bar(pie_labels, category_counts.values, color=pie_colors, edgecolor='white')
        ax2.set_title('目标数量统计', fontsize=13, fontweight='bold')
        ax2.set_ylabel('数量')
        for bar, val in zip(bars, category_counts.values):
            ax2.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3,
                     str(val), ha='center', fontsize=12, fontweight='bold')
        fig2.tight_layout()

        buf2 = io.BytesIO()
        fig2.savefig(buf2, format='png', dpi=150)
        buf2.seek(0)
        plt.close(fig2)

        p_chart2 = doc.add_paragraph()
        p_chart2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img2 = p_chart2.add_run()
        run_img2.add_picture(buf2, width=Inches(4.5))

        # ===== 保存 =====
        if save_path is None:
            save_path = os.path.join(config.RESULTS_DIR, f"{self.session_name}.docx")
        doc.save(save_path)
        return save_path